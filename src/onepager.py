"""
Company Qualitative Research One-Pager — the qual track of a run
(design: docs/feature_qual_onepager.md).

Turns each gate-ACCEPTED company record into an interview-ready kit: research
context, interview brief, respondent survey, semi-structured interview guide,
research priorities. Same pipeline shape as the quant track: the app builds a
prompt from gated data → paste into ChatGPT (or run via src.api_runner) → the
researcher saves `qual/<Brand>_onepager.json` → the qual gate validates it
(repair loop on rejection) → accepted one-pagers render to Markdown.

Anti-hallucination is structural:
  * KNOWN / UNCLEAR lists are computed HERE from the record, not written by
    the LLM;
  * every `basis: fact` claim must cite `source_fields` that exist in the
    record, and may not contain numbers absent from those fields (gate check);
  * hypotheses require `validated_if` — an observable answer that confirms or
    kills them.

Files, per run:
  qual/qual_meta.json          research goal + per-company angle & status
  qual/<Brand>_onepager.json   gated machine artifact
  qual/<Brand>_onepager.md     rendered deliverable
  qual/prompt_qual.md          the current qual prompt to paste
  qual/steps/                  numbered prompt history
"""
from __future__ import annotations

import json
import re
from pathlib import Path

from . import gate
from . import runs

ANGLES = ("competitor", "customer", "partner", "benchmark",
          "market_signal", "acquisition_target")
NEXT_STEPS = ("interview", "desk_research", "expert_call",
              "customer_validation", "partner_outreach", "skip")
SURVEY_TYPES = ("multiple_choice", "ranking", "open")
THEMES = ("market_perception", "customer_problems", "buying_behavior",
          "competitors", "positioning", "product_value", "strengths",
          "weaknesses", "future_trends")
_STD_VALIDATES = ("awareness", "perception", "needs", "differentiation",
                  "buying_criteria")

# provenance markers used in the rendered one-pager
_MARK = {"fact": "Ф", "inference": "В", "hypothesis": "Г"}

# The one-pager's CORE blocks (structure enforced by the qual gate — read-only
# in the Settings tab). Custom blocks live in config/onepager_blocks.yaml.
BUILTIN_BLOCKS = [
    ("context", "Research context: summary + relevance (vs the research goal), "
                "KNOWN/UNCLEAR computed from the record, hypotheses with validated_if."),
    ("interview_brief", "Respondent types (who + why + priority), what we want to "
                        "learn, sensitive topics with approach (avoid/ask_carefully/reframe)."),
    ("survey", "5–8 lightweight questions (≥1 multiple-choice, ≥1 ranking, ≥1 open), "
               "each validating a hypothesis or awareness/perception/needs/"
               "differentiation/buying_criteria."),
    ("interview_guide", "10–15 questions across ≥6 themes: market_perception, "
                        "customer_problems, buying_behavior, competitors, positioning, "
                        "product_value, strengths, weaknesses, future_trends."),
    ("priorities", "Top-3 things to validate (referencing hypotheses), top-3 risks, "
                   "recommended next step (interview/desk_research/expert_call/"
                   "customer_validation/partner_outreach/skip)."),
]

_BLOCKS_FILE = runs.CFG / "onepager_blocks.yaml"


def custom_blocks() -> list[dict]:
    import yaml
    if not _BLOCKS_FILE.exists():
        return []
    return (yaml.safe_load(_BLOCKS_FILE.read_text(encoding="utf-8")) or {}).get("blocks") or []


def _save_blocks(blocks: list[dict]) -> None:
    import yaml
    _BLOCKS_FILE.write_text(
        "# onepager_blocks.yaml — extra one-pager blocks (managed by the app,\n"
        "# tab «3 · Settings» → One-pager layout). Each block is requested from the\n"
        "# researcher in every qual prompt and rendered after the core sections.\n\n"
        + yaml.safe_dump({"blocks": blocks}, allow_unicode=True, sort_keys=False),
        encoding="utf-8")


def add_block(name: str, desc: str) -> None:
    runs._check_col_name(name)
    if name in {b for b, _ in BUILTIN_BLOCKS} or name in {b["name"] for b in custom_blocks()}:
        raise ValueError(f"block '{name}' already exists")
    _save_blocks(custom_blocks() + [{"name": name, "desc": desc.strip() or name}])


def update_block(old_name: str, new_name: str, desc: str) -> None:
    blocks = custom_blocks()
    for b in blocks:
        if b["name"] == old_name:
            if new_name != old_name:
                runs._check_col_name(new_name)
                if new_name in {x for x, _ in BUILTIN_BLOCKS} or \
                   new_name in {x["name"] for x in blocks}:
                    raise ValueError(f"block '{new_name}' already exists")
                b["name"] = new_name
            b["desc"] = desc.strip() or b["desc"]
            _save_blocks(blocks)
            return
    raise ValueError(f"'{old_name}' is not a custom block")


def delete_block(name: str) -> None:
    blocks = [b for b in custom_blocks() if b["name"] != name]
    if len(blocks) == len(custom_blocks()):
        raise ValueError(f"'{name}' is not a custom block")
    _save_blocks(blocks)


def reset_blocks() -> None:
    """Settings «Default»: drop all custom blocks — the report returns to its
    five core sections (future qual prompts only)."""
    _save_blocks([])


# ── qual meta (goal + selection + angles) ─────────────────────────────────────
def qual_dir(run_dir: Path) -> Path:
    d = run_dir / "qual"
    d.mkdir(exist_ok=True)
    return d


def load_meta(run_dir: Path) -> dict:
    f = qual_dir(run_dir) / "qual_meta.json"
    if f.exists():
        return json.loads(f.read_text(encoding="utf-8"))
    return {"research_goal": "", "companies": {}}


def save_meta(run_dir: Path, meta: dict) -> None:
    (qual_dir(run_dir) / "qual_meta.json").write_text(
        json.dumps(meta, ensure_ascii=False, indent=2), encoding="utf-8")


def setup(run_dir: Path, goal: str, angles: dict[str, str],
          manual: dict | None = None) -> dict:
    """Start (or update) the qual track: research goal + selected companies
    with their confirmed angles ({brand: angle}). `manual` carries the
    user-provided context for manually added targets
    ({brand: {"segment": …, "notes": …}}) — run-backed and manual coexist."""
    goal = goal.strip()
    if not goal:
        raise ValueError("research goal is required — it frames relevance, "
                         "hypotheses and the angle for every one-pager")
    bad = [a for a in angles.values() if a not in ANGLES]
    if bad:
        raise ValueError(f"unknown angle(s): {bad}; allowed: {ANGLES}")
    meta = load_meta(run_dir)
    meta["research_goal"] = goal
    for brand, angle in angles.items():
        entry = {**meta["companies"].get(brand, {}), "angle": angle}
        if manual and brand in manual:
            entry["manual"] = {"segment": str(manual[brand].get("segment", "")).strip(),
                               "notes": str(manual[brand].get("notes", "")).strip()}
        meta["companies"][brand] = entry
    save_meta(run_dir, meta)
    return meta


def manual_entry(brand: str, manual: dict) -> dict:
    """Gate-entry-shaped target built ONLY from user-provided context (name,
    segment, optional notes). The synthetic record is explicitly marked
    manual_target and is NEVER a substitute for a gate-verified quantitative
    record — whatever the user did not state stays unclear / inference /
    hypothesis territory, and the segment is context, not a company fact."""
    fields = {}
    if manual.get("segment"):
        fields["segment"] = {"value": str(manual["segment"]).strip(),
                             "source": "user-provided", "confidence": "high"}
    if manual.get("notes"):
        fields["user_notes"] = {"value": str(manual["notes"]).strip(),
                                "source": "user-provided", "confidence": "high"}
    record = {"entity": brand, "manual_target": True, "fields": fields,
              "review_flags": ["manual target — добавлено вручную, без "
                               "количественного исследования за записью"]}
    return {"entity": brand, "stem": runs._slug(brand), "record": record,
            "manual": True}


def target_entries(g: dict, qmeta: dict) -> dict:
    """brand → gate-entry for every selected qual target. A gate-ACCEPTED
    quantitative record always wins over a manual entry with the same
    (normalized) name — that is also the dedup rule; manual targets only fill
    the gaps run data cannot cover."""
    rec_by_entity = {e["entity"]: e for e in g["accepted"]}
    by_norm = {runs._norm(e["entity"]): e for e in g["accepted"]}
    out = {}
    for brand, info in (qmeta.get("companies") or {}).items():
        e = rec_by_entity.get(brand) or by_norm.get(runs._norm(brand))
        if e is None and isinstance(info.get("manual"), dict):
            e = manual_entry(brand, info["manual"])
        if e is not None:
            out[brand] = e
    return out


def remove_target(run_dir: Path, brand: str) -> None:
    """Drop a company (run-backed or manual) from the qual track."""
    meta = load_meta(run_dir)
    if brand in meta.get("companies", {}):
        meta["companies"].pop(brand)
        save_meta(run_dir, meta)


def create_manual_run(name: str = "manual qual targets") -> Path:
    """Standalone container for qualitative research WITHOUT a quantitative
    run: a minimal run dir whose gate has zero records, so every target added
    to it is manual by construction."""
    from datetime import datetime
    run_id = f"{datetime.now().strftime('%Y-%m-%d_%H%M')}_{runs._slug(name)}_qual"
    rd = runs.LOGS / run_id
    (rd / "agent_runs").mkdir(parents=True, exist_ok=True)
    schema = runs.load_schema()
    meta = {"run_id": run_id, "market": name, "depth": "superficial",
            "model": "chatgpt",
            "output_language": schema.get("output_language", "Russian"),
            "geo": schema.get("geo"), "status": "qual",
            "created_at": runs._now(), "xlsx": None}
    (rd / "run.json").write_text(json.dumps(meta, ensure_ascii=False, indent=2),
                                 encoding="utf-8")
    return rd


def propose_angle(rec: dict) -> str:
    """Deterministic default; the analyst confirms/overrides in the app."""
    fields = rec.get("fields") or {}
    em = rec.get("entity_match") or {}
    et = gate.normalize_entity_type(em.get("entity_type") or em.get("type")
                                    or gate.value_of(fields.get("entity_type"))) or ""
    rev = gate.normalize_money(gate.value_of(fields.get("total_revenue_2025")) or "")
    yoy = str(gate.value_of(fields.get("revenue_yoy_24_25")) or "")
    try:
        rev_mln = float(re.sub(r"[^\d,.]", "", (rev or "0").replace(" ", ""))
                        .replace(",", ".") or 0)
    except ValueError:
        rev_mln = 0
    growing = "+" in yoy and not yoy.startswith("-")
    if et == "product" and 0 < rev_mln < 500 and growing:
        return "acquisition_target"
    if rev_mln >= 10_000:
        return "benchmark"
    if et in ("company", "brand", "group", "product") and rev_mln > 0:
        return "competitor"
    return "market_signal"


# ── deterministic KNOWN / UNCLEAR (the trusted part of the one-pager) ─────────
def known_unclear(rec: dict, schema_fields: list[str]) -> tuple[list[dict], list[dict]]:
    fields = rec.get("fields") or {}
    known, unclear = [], []
    for name in schema_fields:
        f = fields.get(name)
        v = gate.value_of(f)
        if v is not None and not gate.is_placeholder(v):
            item = {"field": name, "value": str(v)[:200],
                    "confidence": (f.get("confidence") if isinstance(f, dict) else None) or "medium",
                    "source": gate.domain((f or {}).get("source", "")) if isinstance(f, dict) else ""}
            known.append(item)
            if isinstance(f, dict) and f.get("conflict"):
                unclear.append({"field": name, "reason": "conflict",
                                "note": "A/B sources disagreed"})
            elif item["confidence"] == "low":
                unclear.append({"field": name, "reason": "low_confidence", "note": ""})
        else:
            unclear.append({"field": name, "reason": "blank", "note": ""})
    return known, unclear


def _segment_neighbors(g: dict, entity: str) -> list[dict]:
    """Other accepted companies in the same segment — so questions can name
    real competitors instead of 'your competitors'."""
    def seg(e):
        return str(gate.value_of((e["record"].get("fields") or {}).get("segment")) or "")
    mine = next((seg(e) for e in g["accepted"] if e["entity"] == entity), "")
    out = []
    for e in g["accepted"]:
        if e["entity"] == entity or (mine and seg(e) != mine):
            continue
        f = e["record"].get("fields") or {}
        out.append({"brand": e["entity"],
                    "positioning": str(gate.value_of(f.get("positioning")) or "")[:120],
                    "revenue_2025": gate.normalize_money(
                        gate.value_of(f.get("total_revenue_2025")) or "") or ""})
    return out[:6]


# ── the generation prompt ─────────────────────────────────────────────────────
def build_qual_prompt(run_dir: Path, entries: list[dict], meta_run: dict,
                      qmeta: dict, g: dict) -> str:
    schema_fields = [f["name"] for f in runs.load_schema()["fields"]]
    lang = meta_run["output_language"]
    save = f"logs/{meta_run['run_id']}/qual"
    blocks = []
    for e in entries:
        is_manual = bool(e["record"].get("manual_target"))
        sf = schema_fields + ["user_notes"] if is_manual else schema_fields
        known, unclear = known_unclear(e["record"], sf)
        angle = qmeta["companies"].get(e["entity"], {}).get("angle") or propose_angle(e["record"])
        manual_note = ("""
⚠️ MANUAL TARGET — no verified quantitative record exists for this company.
KNOWN below is ONLY the user-provided context (name, segment, notes). Treat
EVERYTHING else as unclear, inference or hypothesis — never as fact. Do not
invent figures, and do not present the segment or market as company facts.
""" if is_manual else "")
        blocks.append(f"""### 🎯 {e['entity']}  — angle: **{angle}**{manual_note}
Save to: `{save}/{e['stem']}_onepager.json` (`"entity"` must be exactly «{e['entity']}»)

RECORD (the ONLY permitted fact base):
```json
{json.dumps(e['record'], ensure_ascii=False)}
```
KNOWN (computed from the record — you may rephrase, never extend):
```json
{json.dumps(known, ensure_ascii=False)}
```
UNCLEAR (computed — seed «what remains unclear» and your hypotheses from these):
```json
{json.dumps(unclear, ensure_ascii=False)}
```
SEGMENT NEIGHBORS (name them in questions instead of «ваши конкуренты»):
```json
{json.dumps(_segment_neighbors(g, e['entity']), ensure_ascii=False)}
```""")
    companies = "\n\n".join(blocks)
    return f"""# Qualitative research one-pagers — {meta_run['market']} ({len(entries)} companies)

You are a **research designer**, not a researcher. Do NOT browse and do NOT add
facts. Work ONLY from the injected record and market context. Anything you
believe that is not in the record is a **hypothesis** — label it as such.

**Research goal (frame everything against this):** {qmeta['research_goal']}

## Provenance — structural, machine-checked
Every claim carries `basis`: `fact` | `inference` | `hypothesis`.
- `fact` requires `source_fields` naming record fields; its text may not contain
  a number or proper name absent from those fields (auto-rejected otherwise).
- `inference` states its `rationale` resting on facts.
- `hypothesis` requires `validated_if` — what an interview/survey answer must
  show to confirm or kill it. `status` starts as `"untested"`.

## Counts & enums (auto-rejected outside these)
- survey: 5–8 questions, ≥1 `multiple_choice`, ≥1 `ranking`, ≥1 `open`;
  mc/ranking need `options`; each question `validates` an `H*` id or one of
  {', '.join(_STD_VALIDATES)}.
- interview_guide: 10–15 questions across ≥6 themes of: {', '.join(THEMES)};
  each question may `targets` H* ids (they must exist).
- priorities: exactly 3 `validate` (referencing H*) + 3 `risks`;
  `next_step.action` ∈ {' / '.join(NEXT_STEPS)}.
- respondents: ≥2 types, each with `who` (concrete role) and `why`.
- angle is FIXED per company (given above) — design questions for that angle.
- No leading questions; never ask what the record already knows at high
  confidence; every hypothesis is targeted by ≥1 question.
- sensitive: consider private-company finances, layoffs, litigation, sanctions,
  ownership → `approach`: avoid | ask_carefully | reframe.
- Prose in **{lang}**, proper names verbatim.
{_custom_blocks_rules()}

## Output — STRICT JSON per company, saved to the path given above
```json
{{"entity": "…", "angle": "…", "research_goal": "…",
  "context": {{
    "summary":   {{"text": "…", "basis": "fact", "source_fields": ["description", "segment"]}},
    "relevance": {{"text": "…", "basis": "inference", "rationale": "…"}},
    "known":   [{{"text": "…", "source_field": "…", "confidence": "high"}}],
    "unclear": [{{"text": "…", "source_field": "…", "reason": "blank"}}],
    "hypotheses": [{{"id": "H1", "text": "…", "basis": "hypothesis",
                     "grounds": ["…"], "validated_if": "…", "status": "untested"}}]}},
  "interview_brief": {{
    "respondents": [{{"type": "customer", "who": "…", "why": "…", "priority": 1}}],
    "learn": ["…"],
    "sensitive": [{{"topic": "…", "why": "…", "approach": "ask_carefully"}}]}},
  "survey": [{{"id": "S1", "type": "multiple_choice", "text": "…",
               "options": ["…"], "validates": "H1"}}],
  "interview_guide": [{{"theme": "market_perception",
                        "questions": [{{"id": "Q1", "text": "…", "targets": ["H1"]}}]}}],
  "priorities": {{"validate": ["… (H1)"], "risks": ["…"],
                  "next_step": {{"action": "interview", "why": "…"}}}}}}
```

## Companies

{companies}

## When done
One line per company (angle, #hypotheses, recommended next step) and STOP.
"""


def _custom_blocks_rules() -> str:
    blocks = custom_blocks()
    if not blocks:
        return ""
    items = "\n".join(f"- `{b['name']}` — {b['desc']}" for b in blocks)
    return (f"\n## Additional blocks (add a top-level `custom_blocks` object)\n"
            f"For each company also fill `custom_blocks: {{\"<name>\": \"…prose…\"}}` with:\n"
            f"{items}\n"
            f"Same provenance discipline: anything not grounded in the record is a "
            f"labeled guess (\"гипотеза: …\").\n")


# ── qual gate ─────────────────────────────────────────────────────────────────
def validate_onepager(op: dict, rec: dict) -> list[dict]:
    issues: list[dict] = []

    def add(field, severity, code, reason):
        issues.append({"field": field, "severity": severity, "code": code, "reason": reason})

    rec_fields = rec.get("fields") or {}
    rec_digits = {n: re.sub(r"\D", "", str(gate.value_of(f) or ""))
                  for n, f in rec_fields.items()}

    if op.get("angle") not in ANGLES:
        add("angle", "reject", "bad-enum", f"angle «{op.get('angle')}» not in {ANGLES}")

    ctx = op.get("context") or {}
    hyps = ctx.get("hypotheses") or []
    hyp_ids = set()
    for h in hyps:
        hid = h.get("id") or "?"
        hyp_ids.add(hid)
        if not str(h.get("validated_if") or "").strip():
            add(f"hypothesis {hid}", "reject", "untestable-hypothesis",
                "no validated_if — state what an answer must show to confirm or kill it")
    if not hyps:
        add("context.hypotheses", "reject", "counts", "no hypotheses — nothing to validate")

    # fact integrity: fact claims must cite existing fields and add no numbers
    def check_fact(item, where):
        if not isinstance(item, dict) or item.get("basis") != "fact":
            return
        srcs = item.get("source_fields") or ([item["source_field"]] if item.get("source_field") else [])
        missing = [s for s in srcs if s not in rec_fields]
        if not srcs or missing:
            add(where, "reject", "fact-not-in-record",
                f"fact cites {missing or 'no'} source_fields not present in the record")
            return
        allowed = "".join(rec_digits.get(s, "") for s in srcs)
        for num in re.findall(r"\d[\d\s .,]{3,}\d", str(item.get("text", ""))):
            if re.sub(r"\D", "", num) not in allowed:
                add(where, "reject", "fact-not-in-record",
                    f"fact contains number «{num.strip()}» absent from its source_fields")
                return

    check_fact(ctx.get("summary"), "context.summary")
    for i, k in enumerate(ctx.get("known") or []):
        sf = k.get("source_field")
        if sf and sf not in rec_fields:
            add(f"context.known[{i}]", "reject", "fact-not-in-record",
                f"source_field «{sf}» not in the record")

    survey = op.get("survey") or []
    if not 5 <= len(survey) <= 8:
        add("survey", "reject", "counts", f"{len(survey)} questions — must be 5–8")
    types_seen = {q.get("type") for q in survey}
    for t in SURVEY_TYPES:
        if t not in types_seen:
            add("survey", "reject", "counts", f"no {t} question in the mix")
    for q in survey:
        if q.get("type") not in SURVEY_TYPES:
            add(f"survey {q.get('id', '?')}", "reject", "bad-enum",
                f"type «{q.get('type')}» not in {SURVEY_TYPES}")
        if q.get("type") in ("multiple_choice", "ranking") and not q.get("options"):
            add(f"survey {q.get('id', '?')}", "reject", "options-missing",
                "multiple_choice/ranking question without options")
        v = str(q.get("validates") or "")
        if v and v not in _STD_VALIDATES and v not in hyp_ids:
            add(f"survey {q.get('id', '?')}", "reject", "orphan-question",
                f"validates «{v}» — not a hypothesis id or one of {_STD_VALIDATES}")

    guide = op.get("interview_guide") or []
    themes = {t.get("theme") for t in guide}
    n_q = sum(len(t.get("questions") or []) for t in guide)
    if not 10 <= n_q <= 15:
        add("interview_guide", "reject", "counts", f"{n_q} questions — must be 10–15")
    if len(themes & set(THEMES)) < 6:
        add("interview_guide", "reject", "counts",
            f"only {len(themes & set(THEMES))} known themes — need ≥6 of {THEMES}")
    targeted = set()
    for t in guide:
        if t.get("theme") not in THEMES:
            add(f"theme {t.get('theme')}", "reject", "bad-enum",
                f"theme not in {THEMES} — RE-FILE its questions under one of "
                f"those themes (merge into an existing block); do NOT delete "
                f"them, the 10–15 total question count still applies")
        for q in t.get("questions") or []:
            for h in q.get("targets") or []:
                if h not in hyp_ids:
                    add(f"guide {q.get('id', '?')}", "reject", "orphan-question",
                        f"targets unknown hypothesis «{h}»")
                targeted.add(h)
    for q in survey:
        if str(q.get("validates") or "") in hyp_ids:
            targeted.add(q["validates"])
    for hid in sorted(hyp_ids - targeted):
        add(f"hypothesis {hid}", "warn", "orphan-hypothesis",
            "no survey or guide question targets it")

    pri = op.get("priorities") or {}
    if len(pri.get("validate") or []) != 3 or len(pri.get("risks") or []) != 3:
        add("priorities", "reject", "counts", "need exactly 3 validate + 3 risks")
    if (pri.get("next_step") or {}).get("action") not in NEXT_STEPS:
        add("priorities.next_step", "reject", "bad-enum",
            f"action not in {NEXT_STEPS}")

    resp = (op.get("interview_brief") or {}).get("respondents") or []
    if len(resp) < 2:
        add("interview_brief.respondents", "warn", "counts",
            "fewer than 2 respondent types")

    cb = op.get("custom_blocks") or {}
    for b in custom_blocks():
        if not str(cb.get(b["name"]) or "").strip():
            add(f"custom_blocks.{b['name']}", "warn", "custom-block-missing",
                "requested block not filled")

    # placeholders anywhere in the texts
    for path, txt in _walk_texts(op):
        if gate.is_placeholder(txt):
            add(path, "reject", "placeholder", f"placeholder value «{txt}»")
    return issues


def _walk_texts(node, path="op"):
    if isinstance(node, dict):
        for k, v in node.items():
            yield from _walk_texts(v, f"{path}.{k}")
    elif isinstance(node, list):
        for i, v in enumerate(node):
            yield from _walk_texts(v, f"{path}[{i}]")
    elif isinstance(node, str) and node.strip():
        yield path, node


def gate_qual(run_dir: Path) -> dict:
    """Validate all onepager JSONs against their records; render accepted ones."""
    g = runs.run_gate(run_dir, write_report=False)
    qmeta = load_meta(run_dir)
    targets = target_entries(g, qmeta)   # run-backed records win; manual fill gaps
    qd = qual_dir(run_dir)
    out = {"accepted": [], "rejected": [], "pending": [], "records_gate": g}
    for brand, info in qmeta.get("companies", {}).items():
        rec_e = targets.get(brand)
        op_path = _find_onepager(qd, brand)
        if op_path is None or rec_e is None:
            out["pending"].append(brand)
            continue
        try:
            op = json.loads(op_path.read_text(encoding="utf-8"))
            issues = validate_onepager(op, rec_e["record"])
        except Exception as ex:
            op, issues = {}, [{"field": "_file", "severity": "reject",
                               "code": "bad-json", "reason": str(ex)}]
        verdict = "rejected" if any(i["severity"] == "reject" for i in issues) else "accepted"
        entry = {"entity": brand, "path": op_path, "stem": op_path.name[:-len("_onepager.json")],
                 "op": op, "record": rec_e["record"], "issues": issues, "verdict": verdict}
        out[verdict].append(entry)
        info["status"] = verdict
        if verdict == "accepted":
            md = render_md(op, rec_e["record"], qmeta, runs._load_meta(run_dir))
            (op_path.parent / f"{entry['stem']}_onepager.md").write_text(md, encoding="utf-8")
    save_meta(run_dir, qmeta)
    return out


def _find_onepager(qd: Path, brand: str) -> Path | None:
    for p in sorted(qd.glob("*_onepager.json")):
        try:
            if runs._norm(json.loads(p.read_text(encoding="utf-8")).get("entity", "")) == runs._norm(brand):
                return p
        except Exception:
            continue
    slug = runs._slug(brand)
    p = qd / f"{slug}_onepager.json"
    return p if p.exists() else None


# ── state machine (mirrors runs.next_prompt) ──────────────────────────────────
def next_qual_prompt(run_dir: Path, batch: int = 2) -> tuple[str, str]:
    batch = max(1, int(batch))
    qmeta = load_meta(run_dir)
    if not qmeta.get("research_goal") or not qmeta.get("companies"):
        raise SystemExit("Qual track not set up — enter the research goal and "
                         "select companies first (tab «2 · Qualitative research»).")
    meta_run = runs._load_meta(run_dir)
    q = gate_qual(run_dir)
    g = q["records_gate"]
    targets = target_entries(g, qmeta)
    if q["pending"]:
        entries = [targets[b] for b in q["pending"] if b in targets][:batch]
        if not entries:
            raise SystemExit("Selected companies have no gate-accepted records — "
                             "finish the quant repair loop first (or add them as "
                             "manual targets with their own context).")
        kind, text = "qual-research", build_qual_prompt(run_dir, entries, meta_run, qmeta, g)
    elif q["rejected"]:
        kind, text = "qual-repair", _render_qual_repair(meta_run, qmeta, q["rejected"][:batch])
    else:
        kind = "done"
        if report_is_stale(run_dir):
            build_report(run_dir)
        text = (f"# All {len(q['accepted'])} one-pagers passed the qual gate.\n\n"
                f"Final report (executive summary + one-pagers): "
                f"`qual/{report_path(run_dir).name}` — use «Open report».\n")
    _issue_qual_prompt(run_dir, kind, text)
    return kind, text


def _issue_qual_prompt(run_dir: Path, kind: str, text: str) -> None:
    qd = qual_dir(run_dir)
    cur = qd / "prompt_qual.md"
    if cur.exists() and cur.read_text(encoding="utf-8") == text:
        return
    steps = qd / "steps"
    steps.mkdir(exist_ok=True)
    nn = len(list(steps.glob("*.md"))) + 1
    (steps / f"{nn:02d}_{kind}.md").write_text(text, encoding="utf-8")
    cur.write_text(text, encoding="utf-8")
    runs._event(run_dir, "qual_prompt_issued", kind=kind, step=nn)


def _render_qual_repair(meta_run: dict, qmeta: dict, rejected: list[dict]) -> str:
    save = f"logs/{meta_run['run_id']}/qual"
    blocks = []
    for e in rejected:
        items = "\n".join(f"- **{i['field']}** [{i['code']}]: {i['reason']}"
                          for i in e["issues"] if i["severity"] == "reject")
        blocks.append(f"## {e['entity']}  —  `{save}/{e['stem']}_onepager.json`\n{items}")
    body = "\n\n".join(blocks)
    return f"""# Qual repair pass — {meta_run['market']}

The one-pagers below FAILED validation. EDIT each JSON in place: fix ONLY the
listed problems, keep everything else. Same rules as the qual prompt — no new
facts (anything not in the record is a labeled hypothesis with `validated_if`),
counts and enums are re-checked, prose in **{meta_run['output_language']}**.
Research goal: {qmeta['research_goal']}

{body}

## When done
Re-save each file and reply with one line per company saying what changed.
"""


# ── renderer ──────────────────────────────────────────────────────────────────
def render_md(op: dict, rec: dict, qmeta: dict, meta_run: dict) -> str:
    ctx = op.get("context") or {}
    ib = op.get("interview_brief") or {}
    pri = op.get("priorities") or {}

    def mark(item, extra=""):
        b = _MARK.get((item or {}).get("basis", ""), "?")
        return f"**[{b}{extra}]** {item.get('text', '')}"

    known = "\n".join(f"- **[Ф·{k.get('confidence', '?')}]** {k.get('text', '')}"
                      f" `({k.get('source_field', '')})`"
                      for k in ctx.get("known") or [])
    unclear = "\n".join(f"- ❓ {u.get('text', '')} `({u.get('source_field', '')}: "
                        f"{u.get('reason', '')})`" for u in ctx.get("unclear") or [])
    hyps = "\n".join(f"- **[Г·{h.get('status', 'untested')}] {h.get('id')}** — "
                     f"{h.get('text', '')}\n  - ✔️ validated if: {h.get('validated_if', '')}"
                     for h in ctx.get("hypotheses") or [])
    resp = "\n".join(f"{r.get('priority', i + 1)}. **{r.get('type', '')}** — "
                     f"{r.get('who', '')} · {r.get('why', '')}"
                     for i, r in enumerate(ib.get("respondents") or []))
    learn = "\n".join(f"- {x}" for x in ib.get("learn") or [])
    sens = "\n".join(f"- ⚠️ **{s.get('topic', '')}** ({s.get('approach', '')}) — "
                     f"{s.get('why', '')}" for s in ib.get("sensitive") or [])
    survey = "\n".join(
        f"| {q.get('id')} | {q.get('type')} | {q.get('text')} | "
        f"{'; '.join(q.get('options') or []) or '—'} | {q.get('validates', '')} |"
        for q in op.get("survey") or [])
    guide = ""
    qn = 0
    for t in op.get("interview_guide") or []:
        guide += f"\n**{t.get('theme', '')}**\n"
        for q in t.get("questions") or []:
            qn += 1
            tg = f"  _(→ {', '.join(q.get('targets') or [])})_" if q.get("targets") else ""
            guide += f"{qn}. {q.get('text', '')}{tg}\n"
    ns = pri.get("next_step") or {}
    n_known = len(ctx.get("known") or [])
    n_unclear = len(ctx.get("unclear") or [])
    cb = op.get("custom_blocks") or {}
    extra = ""
    if cb:
        extra = "\n## 6 · Additional blocks\n" + "\n".join(
            f"\n**{k}**\n\n{v}\n" for k, v in cb.items())

    manual_mark = (" · ✍ manual target (context provided by the analyst, no "
                   "quantitative research behind it)"
                   if rec.get("manual_target") else "")
    return f"""# {op.get('entity', '?')} — Qualitative Research One-Pager{manual_mark}

**Angle:** {op.get('angle', '?')} · **Next step:** {ns.get('action', '?')} — {ns.get('why', '')}
**Research goal:** {qmeta.get('research_goal', '')}
Легенда: **[Ф]** факт из записи (с confidence) · **[В]** вывод из фактов · **[Г]** гипотеза (проверить)

## 1 · Research context
{mark(ctx.get('summary') or {})}

**Why relevant:** {mark(ctx.get('relevance') or {})}
{('  _' + (ctx.get('relevance') or {}).get('rationale', '') + '_') if (ctx.get('relevance') or {}).get('rationale') else ''}

**Known ({n_known}):**
{known or '—'}

**Unclear ({n_unclear}):**
{unclear or '—'}

**Hypotheses to validate:**
{hyps or '—'}

## 2 · Interview brief
**Respondents:**
{resp or '—'}

**We want to learn:**
{learn or '—'}

**Sensitive topics:**
{sens or '—'}

## 3 · Respondent survey
| # | type | question | options | validates |
|---|---|---|---|---|
{survey}

## 4 · Semi-structured interview guide
{guide or '—'}

## 5 · Research priorities
**Validate (top-3):**
{chr(10).join('- ' + x for x in pri.get('validate') or [])}

**Risks / unknowns (top-3):**
{chr(10).join('- ' + x for x in pri.get('risks') or [])}
{extra}
---
_Data freshness: run `{meta_run.get('run_id', '')}` · {n_known} fields known / {n_unclear} unclear · statuses editable in the JSON (untested → supported / refuted / mixed), re-render via the app._
"""


# ── final report (.docx: executive summary + all one-pagers) ─────────────────
def report_path(run_dir: Path) -> Path:
    meta_run = runs._load_meta(run_dir)
    date = meta_run["run_id"].split("_", 1)[0]
    return qual_dir(run_dir) / f"{date}_{runs._slug(meta_run['market'])}_qual_report.docx"


def report_is_stale(run_dir: Path) -> bool:
    out = report_path(run_dir)
    if not out.exists():
        return True
    newest = max((p.stat().st_mtime
                  for pat in ("*_onepager.json", "*_respondents.json")
                  for p in qual_dir(run_dir).glob(pat)), default=0)
    return newest > out.stat().st_mtime


def build_report(run_dir: Path) -> Path:
    """Merge all ACCEPTED one-pagers into one .docx, executive summary first."""
    from docx import Document
    from docx.shared import Pt

    q = gate_qual(run_dir)
    if not q["accepted"]:
        raise SystemExit("No accepted one-pagers yet — finish the qual track first.")
    meta_run = runs._load_meta(run_dir)
    qmeta = load_meta(run_dir)
    ops = sorted(q["accepted"], key=lambda e: e["entity"].lower())

    doc = Document()
    doc.styles["Normal"].font.size = Pt(10.5)

    def p(text, style=None, bold=False, italic=False):
        par = doc.add_paragraph(style=style)
        run = par.add_run(text)
        run.bold, run.italic = bold, italic
        return par

    doc.add_heading(f"{meta_run['market']} — Qualitative Research Report", 0)
    p(f"Research goal: {qmeta.get('research_goal', '')}", bold=True)
    p(f"Run {meta_run['run_id']} · {len(ops)} companies · "
      f"легенда: [Ф] факт из данных · [В] вывод · [Г] гипотеза (проверить)", italic=True)

    # ── executive summary (compiled from the one-pagers, facts stay facts) ────
    doc.add_heading("Executive summary", 1)
    angles = {}
    for e in ops:
        angles.setdefault(e["op"].get("angle", "?"), []).append(e["entity"])
    p("Охват: " + " · ".join(f"{a}: {', '.join(bs)}" for a, bs in sorted(angles.items())))

    t = doc.add_table(rows=1, cols=5)
    t.style = "Light Grid Accent 1"
    for i, h in enumerate(("Company", "Angle", "Next step", "Key hypothesis", "Known/Unclear")):
        t.rows[0].cells[i].paragraphs[0].add_run(h).bold = True
    for e in ops:
        op = e["op"]
        ctx = op.get("context") or {}
        hyp = (ctx.get("hypotheses") or [{}])[0]
        ns = (op.get("priorities") or {}).get("next_step") or {}
        row = t.add_row().cells
        row[0].text = e["entity"]
        row[1].text = op.get("angle", "")
        row[2].text = ns.get("action", "")
        row[3].text = str(hyp.get("text", ""))[:160]
        row[4].text = f"{len(ctx.get('known') or [])}/{len(ctx.get('unclear') or [])}"

    doc.add_heading("Top validation priorities across the market", 2)
    for e in ops:
        first = ((e["op"].get("priorities") or {}).get("validate") or [""])[0]
        if first:
            p(f"{e['entity']}: {first}", style="List Bullet")
    doc.add_heading("Key risks / unknowns", 2)
    seen = set()
    for e in ops:
        for r in (e["op"].get("priorities") or {}).get("risks") or []:
            k = r.strip().lower()
            if k and k not in seen:
                seen.add(k)
                p(r, style="List Bullet")

    # optional respondent shortlist (only accepted sourcing files; absent when
    # the optional stage was never run)
    try:
        from . import respondents as _resp
        resp_docs = _resp.accepted_docs(run_dir)
    except Exception:
        resp_docs = {}
    if resp_docs.get("market"):
        doc.add_heading("Respondent shortlist — market level", 2)
        p("Публичные профессиональные данные; контакт — только через публичный "
          "профиль. Порядок: приоритет 1 → 3.", italic=True)
        for c in sorted(resp_docs["market"].get("candidates") or [],
                        key=lambda c: (c.get("priority", 9), c.get("name", ""))):
            p(_resp.format_candidate(c), style="List Bullet")

    # ── one company per section ───────────────────────────────────────────────
    for e in ops:
        op, rec = e["op"], e["record"]
        ctx = op.get("context") or {}
        ib = op.get("interview_brief") or {}
        pri = op.get("priorities") or {}
        ns = pri.get("next_step") or {}
        doc.add_page_break()
        doc.add_heading(e["entity"], 1)
        p(f"Angle: {op.get('angle', '')} · Next step: {ns.get('action', '')} — "
          f"{ns.get('why', '')}", italic=True)

        doc.add_heading("1 · Research context", 2)
        p("[Ф] " + str((ctx.get("summary") or {}).get("text", "")))
        p("[В] Why relevant: " + str((ctx.get("relevance") or {}).get("text", "")))
        p("Known:", bold=True)
        for k in ctx.get("known") or []:
            p(f"[Ф·{k.get('confidence', '?')}] {k.get('text', '')}  ({k.get('source_field', '')})",
              style="List Bullet")
        p("Unclear:", bold=True)
        for u in ctx.get("unclear") or []:
            p(f"{u.get('text', '')}  ({u.get('source_field', '')}: {u.get('reason', '')})",
              style="List Bullet")
        p("Hypotheses:", bold=True)
        for h in ctx.get("hypotheses") or []:
            p(f"[Г·{h.get('status', 'untested')}] {h.get('id')}: {h.get('text', '')} — "
              f"validated if: {h.get('validated_if', '')}", style="List Bullet")

        doc.add_heading("2 · Interview brief", 2)
        for r in ib.get("respondents") or []:
            p(f"{r.get('priority', '')}. {r.get('type', '')} — {r.get('who', '')} · "
              f"{r.get('why', '')}", style="List Bullet")
        if ib.get("learn"):
            p("We want to learn: " + "; ".join(ib["learn"]))
        for s in ib.get("sensitive") or []:
            p(f"⚠ {s.get('topic', '')} ({s.get('approach', '')}) — {s.get('why', '')}",
              style="List Bullet")

        doc.add_heading("3 · Respondent survey", 2)
        st = doc.add_table(rows=1, cols=4)
        st.style = "Light Grid Accent 1"
        for i, h in enumerate(("#", "Question", "Type / options", "Validates")):
            st.rows[0].cells[i].paragraphs[0].add_run(h).bold = True
        for qq in op.get("survey") or []:
            row = st.add_row().cells
            row[0].text = str(qq.get("id", ""))
            row[1].text = str(qq.get("text", ""))
            row[2].text = qq.get("type", "") + (
                ": " + "; ".join(qq.get("options") or []) if qq.get("options") else "")
            row[3].text = str(qq.get("validates", ""))

        doc.add_heading("4 · Interview guide", 2)
        n = 0
        for th in op.get("interview_guide") or []:
            p(str(th.get("theme", "")), bold=True)
            for qq in th.get("questions") or []:
                n += 1
                tg = f"  (→ {', '.join(qq.get('targets') or [])})" if qq.get("targets") else ""
                p(f"{n}. {qq.get('text', '')}{tg}", style="List Bullet")

        doc.add_heading("5 · Research priorities", 2)
        p("Validate:", bold=True)
        for x in pri.get("validate") or []:
            p(x, style="List Bullet")
        p("Risks:", bold=True)
        for x in pri.get("risks") or []:
            p(x, style="List Bullet")

        rd_doc = resp_docs.get(e["entity"])
        if rd_doc:
            doc.add_heading("6 · Respondent candidates (public professional data)", 2)
            p("Названные кандидаты под архетипы из раздела 2; контакт — только "
              "через публичный профиль.", italic=True)
            for c in sorted(rd_doc.get("candidates") or [],
                            key=lambda c: (c.get("priority", 9), c.get("name", ""))):
                p(_resp.format_candidate(c), style="List Bullet")

        cb = op.get("custom_blocks") or {}
        if cb:
            doc.add_heading("7 · Additional blocks" if rd_doc else "6 · Additional blocks", 2)
            for k, v in cb.items():
                p(k, bold=True)
                p(str(v))

    out = report_path(run_dir)
    doc.save(out)
    runs._event(run_dir, "qual_report_built", companies=len(ops), path=str(out))
    return out


def progress(run_dir: Path) -> dict:
    qmeta = load_meta(run_dir)
    if not qmeta.get("companies"):
        return {"selected": 0, "accepted": 0, "rejected": 0, "pending": 0,
                "phase": "not started — set goal & companies"}
    q = gate_qual(run_dir)
    a, r, p = len(q["accepted"]), len(q["rejected"]), len(q["pending"])
    phase = (f"researching ({p} to go)" if p else
             f"repair — {r} rejected" if r else "done — one-pagers rendered")
    return {"selected": len(qmeta["companies"]), "accepted": a,
            "rejected": r, "pending": p, "phase": phase}
